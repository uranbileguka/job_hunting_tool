#!/usr/bin/env python3
import json
import os
import re
import shutil
import subprocess
import sys
from pathlib import Path


def clean_filename(text: str) -> str:
    text = (text or "").strip()
    text = re.sub(r"[^A-Za-z0-9\s-]", "", text)
    text = re.sub(r"\s+", "_", text)
    return text or "Unknown"


def apply_placeholders(text: str, payload: dict) -> str:
    out = (text or "").strip()
    replacements = {
        r"\[company[^\]]*\]": payload.get("companyName", ""),
        r"\[date[^\]]*\]": payload.get("date", ""),
        r"\[location[^\]]*\]": payload.get("location", ""),
        r"\[job\s*title[^\]]*\]": payload.get("jobTitle", ""),
    }
    for pattern, value in replacements.items():
        out = re.sub(pattern, str(value or "").strip(), out, flags=re.IGNORECASE)
    return out


def resolve_soffice() -> str:
    env_soffice = os.getenv("SOFFICE_BIN", "").strip()
    if env_soffice and Path(env_soffice).exists():
        return env_soffice

    found = shutil.which("soffice") or shutil.which("libreoffice")
    if found:
        return found

    candidates = [
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        "/Applications/OpenOffice.app/Contents/MacOS/soffice",
        "/opt/homebrew/bin/soffice",
        "/usr/local/bin/soffice",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return candidate
    return ""


def main() -> int:
    try:
        payload = json.loads(sys.stdin.read() or "{}")
    except Exception:
        print(json.dumps({"error": "Invalid JSON payload"}, ensure_ascii=True))
        return 1

    try:
        from docx import Document
        from docx.shared import Pt
    except Exception:
        print(
            json.dumps(
                {
                    "error": "python-docx is not installed. Install from backend/scripts/requirements.txt"
                },
                ensure_ascii=True,
            )
        )
        return 1

    template_path = Path(payload.get("templatePath", "")).expanduser().resolve()
    output_dir = Path(payload.get("outputDir", "")).expanduser().resolve()

    if not template_path.exists():
        print(json.dumps({"error": f"Template file not found: {template_path}"}, ensure_ascii=True))
        return 1

    output_dir.mkdir(parents=True, exist_ok=True)

    date_value = (payload.get("date") or "").strip()
    company = (payload.get("companyName") or "").strip()
    location = (payload.get("location") or "").strip()
    job_title = (payload.get("jobTitle") or "").strip()

    improved = [
        (payload.get("improvedParagraph1") or "").strip(),
        (payload.get("improvedParagraph2") or "").strip(),
        (payload.get("improvedParagraph3") or "").strip(),
        (payload.get("improvedParagraph4") or "").strip(),
        (payload.get("improvedParagraph5") or "").strip(),
    ]
    template_paragraphs = [
        (payload.get("paragraph1") or "").strip(),
        (payload.get("paragraph2") or "").strip(),
        (payload.get("paragraph3") or "").strip(),
        (payload.get("paragraph4") or "").strip(),
        (payload.get("paragraph5") or "").strip(),
    ]

    doc = Document(str(template_path))

    idx_date = 3
    idx_company = 5
    idx_location = 7
    body_slots = [11, 13, 15, 17, 19]

    max_idx = max([idx_date, idx_company, idx_location] + body_slots)
    if len(doc.paragraphs) <= max_idx:
        print(
            json.dumps(
                {
                    "error": "Template format changed. Expected paragraph indexes are missing."
                },
                ensure_ascii=True,
            )
        )
        return 1

    doc.paragraphs[idx_date].text = date_value
    doc.paragraphs[idx_company].text = company
    doc.paragraphs[idx_location].text = location

    for i, p_idx in enumerate(body_slots):
        source_text = ""
        if i < len(improved) and improved[i]:
            source_text = improved[i]
        elif i < len(template_paragraphs):
            source_text = template_paragraphs[i]
        doc.paragraphs[p_idx].text = apply_placeholders(source_text, payload)

    # Keep notebook formatting behavior.
    for p in doc.paragraphs[idx_date:]:
        for run in p.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(11)

    company_clean = clean_filename(company)
    job_clean = clean_filename(job_title)
    filename = f"Uranbileg_CL_{company_clean}_{job_clean}.docx"
    save_path = output_dir / filename
    doc.save(str(save_path))

    pdf_path = save_path.with_suffix(".pdf")
    pdf_created = False
    pdf_error = ""

    # Prefer headless LibreOffice conversion to avoid OS permission prompts.
    soffice = resolve_soffice()
    if soffice:
        try:
            result = subprocess.run(
                [
                    soffice,
                    "--headless",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    str(output_dir),
                    str(save_path),
                ],
                check=True,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
            )
            pdf_created = pdf_path.exists()
            if not pdf_created:
                details = (result.stderr or result.stdout or "").strip()
                pdf_error = (
                    "LibreOffice conversion ran but PDF file was not produced."
                    + (f" Details: {details}" if details else "")
                )
        except Exception as exc:
            pdf_error = f"LibreOffice conversion failed: {exc}"
    elif os.getenv("ENABLE_DOCX2PDF", "").strip() == "1":
        # Optional fallback (may trigger OS permission prompt on macOS).
        try:
            from docx2pdf import convert

            convert(str(save_path), str(pdf_path))
            pdf_created = pdf_path.exists()
        except Exception as exc:
            pdf_error = str(exc)
    else:
        pdf_error = (
            "PDF skipped to avoid permission prompts. Install LibreOffice "
            "(soffice) for headless PDF export, or set ENABLE_DOCX2PDF=1."
        )

    print(
        json.dumps(
            {
                "docxPath": str(save_path),
                "pdfPath": str(pdf_path) if pdf_created else "",
                "pdfCreated": pdf_created,
                "pdfError": pdf_error,
            },
            ensure_ascii=True,
        )
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
