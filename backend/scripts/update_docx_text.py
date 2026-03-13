#!/usr/bin/env python3
import json
import re
import sys
from pathlib import Path


def split_paragraphs(text: str):
    normalized = (text or "").replace("\r\n", "\n").replace("\r", "\n")
    parts = [re.sub(r"\s+", " ", p).strip() for p in normalized.split("\n\n")]
    return [p for p in parts if p]


def iter_cell_paragraphs(cell, visited_cells):
    cell_key = id(cell._tc)
    if cell_key in visited_cells:
        return
    visited_cells.add(cell_key)

    for paragraph in cell.paragraphs:
        yield paragraph

    for table in cell.tables:
        for row in table.rows:
            for nested_cell in row.cells:
                yield from iter_cell_paragraphs(nested_cell, visited_cells)


def iter_table_paragraphs(table, visited_cells):
    for row in table.rows:
        for cell in row.cells:
            yield from iter_cell_paragraphs(cell, visited_cells)


def iter_all_paragraphs(doc):
    visited_cells = set()

    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        yield from iter_table_paragraphs(table, visited_cells)

    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            yield paragraph
        for table in section.header.tables:
            yield from iter_table_paragraphs(table, visited_cells)

        for paragraph in section.footer.paragraphs:
            yield paragraph
        for table in section.footer.tables:
            yield from iter_table_paragraphs(table, visited_cells)


def set_paragraph_text_preserve_format(paragraph, new_text: str):
    runs = list(paragraph.runs)
    if not runs:
        paragraph.add_run(new_text)
        return

    # Keep paragraph/run formatting by reusing existing runs.
    runs[0].text = new_text
    for run in runs[1:]:
        run.text = ""


def main() -> int:
    try:
      payload = json.loads(sys.stdin.read() or "{}")
    except Exception:
      print(json.dumps({"error": "Invalid JSON payload"}, ensure_ascii=True))
      return 1

    file_path = Path(str(payload.get("filePath") or "")).expanduser().resolve()
    text = str(payload.get("text") or "")

    if not file_path.exists():
      print(json.dumps({"error": f"File not found: {file_path}"}, ensure_ascii=True))
      return 1
    if file_path.suffix.lower() != ".docx":
      print(json.dumps({"error": "Only .docx files are supported."}, ensure_ascii=True))
      return 1

    try:
      from docx import Document
    except Exception:
      print(
        json.dumps(
          {"error": "python-docx is not installed. Install from backend/scripts/requirements.txt"},
          ensure_ascii=True,
        )
      )
      return 1

    paragraphs = split_paragraphs(text)
    doc = Document(str(file_path))

    all_paragraphs = list(iter_all_paragraphs(doc))
    existing_count = len(all_paragraphs)
    min_count = min(existing_count, len(paragraphs))

    for i in range(min_count):
      set_paragraph_text_preserve_format(all_paragraphs[i], paragraphs[i])

    # Do not add/remove structural blocks. This preserves original document layout.
    if existing_count > len(paragraphs):
      for i in range(len(paragraphs), existing_count):
        set_paragraph_text_preserve_format(all_paragraphs[i], "")

    doc.save(str(file_path))
    print(json.dumps({"updated": True, "filePath": str(file_path)}, ensure_ascii=True))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
